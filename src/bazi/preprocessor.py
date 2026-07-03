"""预处理层: 格式统一 + 内容清洗 + 结构提取 + 语义分块

上游输入: 原始文件路径 (docx/pdf/doc/txt/epub)
下游输出: TextBlock 列表，供 Pipeline 进入向量化和分类阶段

架构:
    FormatBridge    → Markdown 文本
    ContentCleaner  → 净化的 Markdown
    StructureParser → [(section_path, paragraphs)]
    SemanticChunker → [TextBlock]
"""

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from .config import config


@dataclass
class TextBlock:
    """语义分块产出单元"""
    block_id: str                          # {book_id}_{index:04d}
    text: str                              # 块文本
    section: Optional[str] = None          # 所属章节路径
    genre_tags: list = field(default_factory=list)     # 主流派标签（单个）
    topic_tags: list = field(default_factory=list)     # 主专题标签（单个）
    source: str = ""                       # 来源书名
    secondary_genres: list = field(default_factory=list)  # 次流派
    secondary_topics: list = field(default_factory=list)  # 次专题
    source_file: str = ""                  # 来源文件路径


class FormatBridge:
    """格式统一层: 将各类文档转换为干净的 Markdown 文本

    支持: docx, doc, pdf (含扫描版OCR兜底), txt, epub
    """

    def __init__(self):
        self._md_converter = None
        self._ocr = None

    def _get_md(self):
        if self._md_converter is None:
            try:
                from markitdown import MarkItDown
                self._md_converter = MarkItDown()
            except ImportError:
                self._md_converter = False
        return self._md_converter if self._md_converter is not False else None

    def _get_ocr(self):
        if self._ocr is None:
            try:
                from paddleocr import PaddleOCR
                self._ocr = PaddleOCR(lang=config.ocr_lang, show_log=False)
            except Exception:
                self._ocr = False
        return self._ocr if self._ocr is not False else None

    def convert(self, file_path: str, suffix: str) -> Optional[str]:
        """将文件转换为 Markdown 文本。返回 None 表示无法处理"""
        path = Path(file_path)
        if not path.exists():
            return None

        suffix = suffix.lower().lstrip(".")

        if suffix == "txt":
            return self._convert_txt(path)
        elif suffix == "docx":
            return self._convert_docx(path)
        elif suffix == "doc":
            return self._convert_doc(path)
        elif suffix == "pdf":
            return self._convert_pdf(path)
        elif suffix == "epub":
            return self._convert_epub(path)
        else:
            return None

    def _convert_txt(self, path: Path) -> str:
        try:
            import chardet
            raw = path.read_bytes()
            detected = chardet.detect(raw)
            encoding = detected.get("encoding", "utf-8") or "utf-8"
            return raw.decode(encoding, errors="replace")
        except Exception:
            return path.read_text(encoding="utf-8", errors="replace")

    def _convert_docx(self, path: Path) -> Optional[str]:
        md = self._get_md()
        if md is not None:
            try:
                result = md.convert(str(path))
                if result and result.text_content:
                    return result.text_content
            except Exception:
                pass
        # 兜底: python-docx
        try:
            from docx import Document
            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception:
            return None

    def _convert_doc(self, path: Path) -> Optional[str]:
        md = self._get_md()
        if md is not None:
            try:
                result = md.convert(str(path))
                if result and result.text_content:
                    return result.text_content
            except Exception:
                pass
        # 兜底: 先试 docx 兼容
        try:
            from docx import Document
            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception:
            pass
        return None

    def _convert_pdf(self, path: Path) -> Optional[str]:
        # 尝试提取文本层
        md = self._get_md()
        text_content = None
        if md is not None:
            try:
                result = md.convert(str(path))
                if result and result.text_content:
                    text_content = result.text_content
            except Exception:
                pass

        # 用 pypdf 验证文本层是否为空
        import pypdf
        try:
            reader = pypdf.PdfReader(str(path))
            native_text = "\n".join(
                page.extract_text() or "" for page in reader.pages
            ).strip()
        except Exception:
            native_text = ""

        # 有文本层则返回
        has_text = text_content and len(text_content.strip()) > 100
        if has_text:
            return text_content

        # 文本层为空 → 扫描版PDF → OCR
        if config.ocr_enabled:
            return self._ocr_pdf(path)

        return text_content or native_text or None

    def _ocr_pdf(self, path: Path) -> Optional[str]:
        """PaddleOCR 逐页识别扫描版PDF"""
        ocr = self._get_ocr()
        if ocr is False:
            return None

        # OCR结果缓存
        cache_dir = config.data_dir / "ocr_cache"
        cache_dir.mkdir(parents=True, exist_ok=True)
        file_hash = hashlib.md5(str(path).encode()).hexdigest()
        cache_file = cache_dir / f"{file_hash}.txt"
        if cache_file.exists():
            return cache_file.read_text(encoding="utf-8")

        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            page_count = len(reader.pages)
        except Exception:
            return None

        all_texts = []
        for i in range(page_count):
            try:
                result = ocr.ocr(str(path), cls=False)
                if result and result[0]:
                    page_text = " ".join(
                        line[1][0] for group in result for line in group
                        if line and len(line) > 1
                    )
                    all_texts.append(page_text)
            except Exception:
                continue

        combined = "\n".join(all_texts).strip()
        if combined:
            cache_file.write_text(combined, encoding="utf-8")
        return combined or None

    def _convert_epub(self, path: Path) -> Optional[str]:
        md = self._get_md()
        if md is not None:
            try:
                result = md.convert(str(path))
                if result and result.text_content:
                    return result.text_content
            except Exception:
                pass
        return None


class StructureParser:
    """结构提取层: 解析 Markdown 标题层级，提取章节边界"""

    @staticmethod
    def parse(markdown_text: str) -> list[tuple[Optional[str], list[str]]]:
        """解析 Markdown，返回 [(section_path, [paragraphs]), ...]"""
        if not markdown_text.strip():
            return [(None, [])]

        lines = markdown_text.split("\n")
        sections = []
        current_path = None
        current_paras = []

        # 当前各级标题
        h1 = h2 = h3 = None

        for line in lines:
            stripped = line.strip()

            # H1: #
            m1 = re.match(r"^#\s+(.+)$", stripped)
            if m1:
                if current_paras:
                    sections.append((current_path, current_paras))
                h1 = m1.group(1).strip()
                h2 = h3 = None
                current_path = h1
                current_paras = []
                continue

            # H2: ##
            m2 = re.match(r"^##\s+(.+)$", stripped)
            if m2:
                if current_paras:
                    sections.append((current_path, current_paras))
                h2 = m2.group(1).strip()
                h3 = None
                parts = [p for p in [h1, h2] if p]
                current_path = " / ".join(parts) if parts else None
                current_paras = []
                continue

            # H3: ###
            m3 = re.match(r"^###\s+(.+)$", stripped)
            if m3:
                if current_paras:
                    sections.append((current_path, current_paras))
                h3 = m3.group(1).strip()
                parts = [p for p in [h1, h2, h3] if p]
                current_path = " / ".join(parts) if parts else None
                current_paras = []
                continue

            # 普通段落
            if stripped:
                current_paras.append(stripped)

        # 最后一个section
        if current_path is not None or current_paras:
            sections.append((current_path, current_paras))

        # 如果没有任何section（无标题文档），整篇作为一个section
        if not sections:
            sections.append((None, current_paras))

        return sections


class SemanticChunker:
    """语义分块层: 在章节边界内按自然段 + 句号兜底进行分块

    优先级: 章节边界 > 自然段 > 句号拆分
    """

    @classmethod
    def chunk(cls, sections: list[tuple[Optional[str], list[str]]],
              book_id: str, source_book: str) -> list[TextBlock]:
        """对解析后的章节进行语义分块"""
        blocks = []
        block_index = 0

        for section_path, paragraphs in sections:
            if not paragraphs:
                continue

            current_texts = []
            current_len = 0

            for para in paragraphs:
                para_len = len(para)

                # 长段拆分: > config.chunk_max_size
                if para_len > config.chunk_max_size:
                    # 先累积当前
                    if current_texts:
                        blocks.append(cls._make_block(
                            current_texts, book_id, block_index, section_path, source_book))
                        block_index += 1
                        current_texts = []
                        current_len = 0

                    # 按句号拆分长段
                    sub_blocks = cls._split_long_para(para)
                    for sub in sub_blocks:
                        blocks.append(cls._make_block(
                            [sub], book_id, block_index, section_path, source_book))
                        block_index += 1
                    continue

                # 加入当前段落会导致超标
                if current_len + para_len > config.chunk_max_size and current_texts:
                    blocks.append(cls._make_block(
                        current_texts, book_id, block_index, section_path, source_book))
                    block_index += 1
                    current_texts = [para]
                    current_len = para_len
                else:
                    current_texts.append(para)
                    current_len += para_len

            # 章节末尾，输出剩余
            if current_texts:
                # 短段：如果不足 chunk_min_size 且不是最后章节，留到与下一章节合并
                # 简化：直接输出
                blocks.append(cls._make_block(
                    current_texts, book_id, block_index, section_path, source_book))
                block_index += 1

        return blocks

    @staticmethod
    def _make_block(texts: list[str], book_id: str, idx: int,
                    section: Optional[str], source: str) -> TextBlock:
        return TextBlock(
            block_id=f"{book_id}_{idx:04d}",
            text="\n\n".join(texts),
            section=section,
            source=source,
        )

    @staticmethod
    def _split_long_para(para: str) -> list[str]:
        """将超长段落按句末标点拆分"""
        # 按 。! ？ 切分
        sentences = re.split(r"(?<=[。！？])", para)
        chunks = []
        current = ""

        for sent in sentences:
            if len(current) + len(sent) <= config.chunk_max_size:
                current += sent
            else:
                if current:
                    chunks.append(current)
                # 如果单个句段还是太长，强制截断
                current = sent
                while len(current) > config.chunk_max_size:
                    chunks.append(current[:config.chunk_max_size])
                    current = current[config.chunk_max_size:]

        if current:
            chunks.append(current)

        return chunks


class LLMNormalizer:
    """LLM驱动的文本规范化与去广告标注

    在分块后、向量化前运行。对每个块：
    1. 修复断裂词（"结 婚"→"结婚"）
    2. 合并被误断的短行
    3. 标注垃圾行（广告/水印/页码）用 [AD] 前缀

    仅对疑似有问题的块调用LLM，干净块直接跳过。
    """

    NORMALIZE_PROMPT = """你是一个文本预处理助手。请对以下八字命理文本块做两件事：

1. 【文本规范化】
   - 修复因OCR/PDF提取造成的断裂词（如"结 婚"→"结婚"、"分 析日 干"→"分析日干"）
   - 合并因强制换行被误断的短行
   - 保持原文内容不变，不做删改

2. 【垃圾标记】
   - 标记广告/水印行（如"加微信XXX"、"公众号XXX"、"好又全资源网"等）
   - 标记页眉页脚/页码残留（如孤立的数字"220"、"221"）
   - 用 [AD] 前缀标记每条垃圾行
   - 正文中偶然出现的数字或URL不要标记

输出格式：直接输出处理后的文本，垃圾行前加 [AD] 前缀。

待处理文本：
{block_text}"""

    # 检测哪些块需要规范化
    _BROKEN_WORD_RE = re.compile(r"[\u4e00-\u9fff]\s+[\u4e00-\u9fff]")
    _AD_KEYWORD_RE = re.compile(
        r"微信|公众号|扫码|免费领取|A{0,2}AKK|QQ群|"
        r"添加微信|微信号|www\.|http[s]?://|好又全|"
        r"关注.*公众号|获取更多.*资料|送.*电子书|资源更新|"
        r"加微信|九鼎学堂",
        re.IGNORECASE,
    )

    def __init__(self):
        from .config import config as cfg
        self._cfg = cfg
        self._batch_size = 20

    def _needs_normalization(self, block: TextBlock) -> bool:
        """快速检测：块是否需要LLM规范化"""
        text = block.text
        if not text:
            return False

        # 1. 检测中文空格断裂
        if self._BROKEN_WORD_RE.search(text):
            return True

        # 2. 检测广告关键词
        if self._AD_KEYWORD_RE.search(text):
            return True

        # 3. 检测过短行（可能因PDF换行被误断）
        lines = text.split("\n")
        short_lines = [l for l in lines if 0 < len(l.strip()) < 10]
        if len(short_lines) >= 3:
            return True

        return False

    def _call_llm(self, prompt: str, max_tokens: int = 2048) -> str:
        """调用 Ollama 兼容的 LLM API"""
        from openai import OpenAI

        client = OpenAI(
            base_url=self._cfg.llm_base_url,
            api_key="ollama",
            timeout=self._cfg.classifier_timeout,
            max_retries=self._cfg.classifier_max_retries,
        )
        response = client.chat.completions.create(
            model=self._cfg.llm_model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=max_tokens,
        )
        return response.choices[0].message.content.strip()

    def _normalize_single(self, block: TextBlock) -> TextBlock:
        """规范化单个文本块"""
        prompt = self.NORMALIZE_PROMPT.format(block_text=block.text[:3000])
        try:
            result_text = self._call_llm(prompt, max_tokens=4096)
        except Exception as e:
            # LLM调用失败，回退：原样保留
            print(f"[Normalizer] LLM调用失败 ({block.block_id}): {e}")
            return block

        return TextBlock(
            block_id=block.block_id,
            text=result_text,
            section=block.section,
            source=block.source,
            secondary_genres=block.secondary_genres,
            secondary_topics=block.secondary_topics,
            source_file=block.source_file,
        )

    def _normalize_batch(self, blocks: list[TextBlock]) -> list[TextBlock]:
        """批处理规范化（逐个调用，带进度）"""
        results = []
        for i, block in enumerate(blocks):
            result = self._normalize_single(block)
            results.append(result)
            if (i + 1) % 5 == 0 or len(blocks) <= 5:
                print(f"  [规范化] {i+1}/{len(blocks)}", flush=True)
        return results

    def normalize(self, blocks: list[TextBlock]) -> list[TextBlock]:
        """公共入口：对块列表做选择性规范化

        干净的块原样通过，脏的块调用LLM规范化。
        """
        if not blocks:
            return []

        # 分离干净块和脏块
        clean_blocks = []
        dirty_blocks = []
        for block in blocks:
            if self._needs_normalization(block):
                dirty_blocks.append(block)
            else:
                clean_blocks.append(block)

        if not dirty_blocks:
            return clean_blocks

        # 分批处理脏块
        normalized = []
        for i in range(0, len(dirty_blocks), self._batch_size):
            batch = dirty_blocks[i:i + self._batch_size]
            print(f"[规范化] 批次 {i // self._batch_size + 1}: {len(batch)} 块",
                  flush=True)
            normalized.extend(self._normalize_batch(batch))

        # 合并：保持原顺序
        result = []
        dirty_iter = iter(normalized)
        for block in blocks:
            if block in dirty_blocks:
                result.append(next(dirty_iter))
            else:
                result.append(block)

        return result
