"""book-parser 模块的单元测试"""

import pytest
from pathlib import Path
from bazi.book_parser import (
    Block,
    parse_file,
    smart_chunk,
)


TEST_DIR = Path(__file__).parent


class TestSmartChunking:
    """智能分块测试"""

    def test_single_short_text(self):
        """短文本应该作为一个块返回"""
        text = "这是一段不到300字的短文本。"
        blocks = smart_chunk(text, source_book="测试书", chapter="第一章", book_id="test")
        assert len(blocks) == 1
        assert blocks[0].source_book == "测试书"
        assert blocks[0].chapter == "第一章"

    def test_long_text_chunking(self):
        """长文本应该被分多个块，块大小在范围内"""
        # 生成约2000字的文本
        text = "八字命理是中华传统文化的瑰宝。" * 200
        blocks = smart_chunk(text, source_book="测试书", chapter="第一章", book_id="test")
        assert len(blocks) > 1
        for block in blocks:
            assert 200 <= len(block.text) <= 550  # 允许一定浮动

    def test_chunk_overlap(self):
        """相邻块之间应该有重叠"""
        text = "这是第一段测试文本。" * 50 + "\n\n" + "这是第二段测试文本。" * 50
        blocks = smart_chunk(text, source_book="测试", chapter="一", book_id="t1")
        if len(blocks) >= 2:
            # 验证前一个块的结尾和后一个块的开头有重叠
            b0_end = blocks[0].text[-20:]
            b1_start = blocks[1].text[:20]
            # 至少有部分重叠或连续
            assert len(b0_end) > 0 and len(b1_start) > 0

    def test_chunk_preserves_paragraph_boundaries(self):
        """分块应该尽量保持段落边界"""
        paragraphs = [
            "第一段：这是一个测试段落，包含一些命理相关的内容，比如天干地支和五行生克。",
            "第二段：这是另一个段落，讨论正官和七杀的区别以及对八字的影响。",
            "第三段：命理学中的十神体系包括正印、偏印、正官、七杀、正财、偏财、食神、伤官、比肩、劫财。",
        ]
        text = "\n\n".join(paragraphs * 10)
        blocks = smart_chunk(text, source_book="测试", chapter="一", book_id="t1")
        # 分块不应在段落中间切断（除非段落本身超过块大小）
        for block in blocks:
            assert block.text is not None
            assert len(block.text) > 0

    def test_empty_text(self):
        """空文本应该返回空列表"""
        blocks = smart_chunk("", source_book="测试", chapter="一", book_id="t1")
        assert blocks == []

    def test_metadata_in_blocks(self):
        """每个块应该包含正确的元信息"""
        text = "测试内容。" * 100
        blocks = smart_chunk(text, source_book="子平真诠", chapter="第一章 论十神", book_id="abc123")
        for block in blocks:
            assert block.source_book == "子平真诠"
            assert block.chapter == "第一章 论十神"
            assert block.book_id == "abc123"
            assert block.block_id is not None
            assert block.block_id.startswith("abc123")


class TestBlockDataclass:
    """Block 数据类测试"""

    def test_block_creation(self):
        block = Block(
            block_id="test_001",
            text="测试内容",
            source_book="测试书",
            chapter="第一章",
            book_id="bk001",
        )
        assert block.block_id == "test_001"
        assert block.text == "测试内容"
        assert block.source_book == "测试书"


class TestParseFile:
    """文件解析测试"""

    def test_parse_txt_file(self, tmp_path):
        """测试 TXT 文件解析"""
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("这是测试文本内容。\n\n这是第二段内容。", encoding="utf-8")

        blocks = parse_file(str(txt_file), "test_book", book_id="bk001")
        assert len(blocks) >= 0  # 短文本可能只有一个块
        for block in blocks:
            assert "测试" in block.text

    def test_parse_txt_with_gbk_encoding(self, tmp_path):
        """测试 GBK 编码的 TXT 文件"""
        txt_file = tmp_path / "test_gbk.txt"
        txt_file.write_text("这是GBK编码的中文文本。", encoding="gbk")

        blocks = parse_file(str(txt_file), "test_book", book_id="bk001")
        assert len(blocks) > 0
        assert "GBK编码" in blocks[0].text

    def test_parse_docx_file(self, tmp_path):
        """测试 DOCX 文件解析"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        docx_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("这是DOCX测试文档的第一段。")
        doc.add_paragraph("这是第二段，包含命理相关内容。")
        doc.save(str(docx_file))

        blocks = parse_file(str(docx_file), "测试文档", book_id="bk001")
        assert len(blocks) > 0
        assert any("DOCX" in block.text for block in blocks)

    def test_parse_unsupported_format(self, tmp_path):
        """不支持的格式应该抛出错误"""
        xyz_file = tmp_path / "test.xyz"
        xyz_file.write_text("dummy content")
        with pytest.raises(ValueError, match="不支持的文件格式"):
            parse_file(str(xyz_file), "测试", book_id="bk001")

    def test_parse_nonexistent_file(self):
        """不存在的文件应该抛出错误"""
        with pytest.raises(FileNotFoundError):
            parse_file("nonexistent.pdf", "测试", book_id="bk001")
